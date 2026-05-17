from pathlib import Path
from datetime import datetime


def test_parse_document_path_inputs_extracts_multiple_unique_paths():
    from investment_system.micro_programs.word_doc_to_md import parse_document_path_inputs

    paths = parse_document_path_inputs(
        [
            '"D:\\docs\\memo one.docx"',
            "另一个：D:\\docs\\memo two.pdf",
            "D:\\docs\\memo three.txt",
            "D:\\docs\\old memo.doc",
            "D:\\docs\\memo one.docx",
            "无关：https://example.com/file.pdf",
        ]
    )

    assert paths == [
        Path("D:\\docs\\memo one.docx"),
        Path("D:\\docs\\memo two.pdf"),
        Path("D:\\docs\\memo three.txt"),
        Path("D:\\docs\\old memo.doc"),
    ]


def test_docx_to_markdown_keeps_headings_lists_marks_and_color(tmp_path):
    from docx import Document
    from docx.shared import RGBColor

    from investment_system.micro_programs.word_doc_to_md import docx_to_markdown

    source = tmp_path / "meeting.docx"
    document = Document()
    document.core_properties.title = "测试会议"
    document.add_heading("一级标题", level=1)
    document.add_paragraph("第一条", style="List Bullet")
    document.add_paragraph("第二层", style="List Bullet 2")
    paragraph = document.add_paragraph()
    bold_run = paragraph.add_run("重点")
    bold_run.bold = True
    paragraph.add_run(" 和 ")
    color_run = paragraph.add_run("红色")
    color_run.font.color.rgb = RGBColor(255, 0, 0)
    document.save(source)

    converted = docx_to_markdown(source)

    assert converted.title == "测试会议"
    assert "# 一级标题" in converted.markdown
    assert "- 第一条" in converted.markdown
    assert "  - 第二层" in converted.markdown
    assert "**重点**" in converted.markdown
    assert '<span style="color:#FF0000">红色</span>' in converted.markdown


def test_txt_to_markdown_reads_common_chinese_encoding(tmp_path):
    from investment_system.micro_programs.word_doc_to_md import txt_to_markdown

    source = tmp_path / "memo.txt"
    source.write_bytes("第一行\n第二行".encode("gb18030"))

    converted = txt_to_markdown(source)

    assert converted.title == "memo"
    assert "第一行" in converted.markdown
    assert "第二行" in converted.markdown


def test_pdf_to_markdown_extracts_text(tmp_path):
    import pytest

    fitz = pytest.importorskip("fitz")
    pytest.importorskip("pdfplumber")

    source = tmp_path / "memo.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "PDF memo text")
    pdf.save(source)
    pdf.close()

    from investment_system.micro_programs.word_doc_to_md import pdf_to_markdown

    converted = pdf_to_markdown(source)

    assert converted.title == "memo"
    assert "PDF memo text" in converted.markdown


def test_standardized_markdown_stem_matches_clipboard_rename_rules(tmp_path):
    from investment_system.micro_programs.word_doc_to_md import standardized_markdown_stem

    now = datetime(2026, 5, 17)

    assert standardized_markdown_stem(tmp_path / "20260515_赛维时代策略会.docx", now) == "20260515_赛维时代策略会"
    assert standardized_markdown_stem(tmp_path / "2026-05-15 赛维时代策略会.docx", now) == "20260515_赛维时代策略会"
    assert standardized_markdown_stem(tmp_path / "赛维时代策略会20260515.docx", now) == "20260515_赛维时代策略会"
    assert standardized_markdown_stem(tmp_path / "赛维时代策略会 2026-05-15.docx", now) == "20260515_赛维时代策略会"
    assert standardized_markdown_stem(tmp_path / "251102_公司交流.docx", now) == "20251102_公司交流"
    assert standardized_markdown_stem(tmp_path / "202605_月度纪要.pdf", now) == "202605_月度纪要"


def test_standardized_markdown_stem_uses_mtime_when_no_date(tmp_path):
    from investment_system.micro_programs.word_doc_to_md import standardized_markdown_stem

    source = tmp_path / "无日期纪要.txt"
    source.write_text("content", encoding="utf-8")
    mtime = datetime(2026, 4, 3, 12, 0).timestamp()
    source.touch()
    import os

    os.utime(source, (mtime, mtime))

    assert standardized_markdown_stem(source) == "20260403_无日期纪要"


def test_build_word_markdown_uses_memo_template_and_demotes_body_h1(tmp_path):
    from investment_system.micro_programs.word_doc_to_md import WordDocument, build_word_markdown

    document = WordDocument(
        title="测试会议",
        markdown="# 原标题\n\n**重点**正文",
        source_path=tmp_path / "测试会议.docx",
    )

    content = build_word_markdown(document)

    assert content.startswith("# 基本信息\n")
    assert "- **原文链接**: [测试会议](" in content
    assert "- **行业**: " in content
    assert "# 正文\n\n## 原标题" in content
    assert "# AI 评价" not in content
