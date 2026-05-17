#!/usr/bin/env python3
"""Convert local Word documents into the memo Inbox Markdown template."""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import tempfile
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Iterable


if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

try:
    from investment_system.common.markdown_utils import bump_headings_when_h1_exists, normalize_markdown_output
    from investment_system.common.utils.paths import MEMO_BASE_DIR

    DEFAULT_OUT_DIR = Path(MEMO_BASE_DIR) / "0-Inbox"
except Exception:
    from investment_system.common.markdown_utils import bump_headings_when_h1_exists, normalize_markdown_output

    DEFAULT_OUT_DIR = Path(r"D:\softwares\Obsidian\MyNotes\信息收集器\C会议纪要\0-Inbox")


class WordConvertError(RuntimeError):
    """Raised when a local document cannot be converted."""


@dataclass
class WordDocument:
    title: str
    markdown: str
    source_path: Path


SUPPORTED_SUFFIXES = {".docx", ".doc", ".pdf", ".txt"}
SUPPORTED_SUFFIX_PATTERN = r"(?:docx|doc|pdf|txt)"


def safe_filename(name: str, default: str = "Word文档") -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', " ", name).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    cleaned = cleaned.rstrip(". ")
    return (cleaned or default)[:120]


def expand_two_digit_year(year2: str, now: datetime | None = None) -> str:
    current_year = (now or datetime.now()).year
    century = (current_year // 100) * 100
    return str(century + int(year2))


def parse_pure_number_date(value: str, now: datetime | None = None) -> dict[str, str | None] | None:
    length = len(value)
    if length == 8:
        year, month, day = value[:4], value[4:6], value[6:8]
        if "01" <= month <= "12" and "01" <= day <= "31":
            return {"year": year, "month": month, "day": day, "raw": value}
        return None
    if length == 6:
        yy, mm, dd = value[:2], value[2:4], value[4:6]
        if "01" <= mm <= "12" and "01" <= dd <= "31":
            return {"year": expand_two_digit_year(yy, now), "month": mm, "day": dd, "raw": value}
        year4, month2 = value[:4], value[4:6]
        if "01" <= month2 <= "12" and 1900 <= int(year4) <= 2099:
            return {"year": year4, "month": month2, "day": None, "raw": value}
        return None
    if length == 4:
        yy, mm = value[:2], value[2:4]
        if "01" <= mm <= "12":
            return {"year": expand_two_digit_year(yy, now), "month": mm, "day": None, "raw": value}
    return None


def extract_filename_date(text: str, *, must_start: bool = False, now: datetime | None = None) -> dict[str, str | None] | None:
    if must_start:
        match = re.match(r"^(\d{4,8})", text)
        if match:
            parsed = parse_pure_number_date(match.group(1), now)
            if parsed:
                return parsed
    else:
        for match in re.finditer(r"\d{4,8}", text):
            parsed = parse_pure_number_date(match.group(0), now)
            if parsed:
                return parsed

    patterns: list[tuple[str, bool]] = [
        (r"(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})", True),
        (r"(\d{4})[.\-/](\d{1,2})", False),
        (r"(\d{4})-(\d{2})(\d{2})", True),
        (r"(\d{4})年(\d{1,2})月(\d{1,2})日?", True),
        (r"(\d{4})年(\d{1,2})月", False),
    ]
    for pattern, has_day in patterns:
        regex = "^" + pattern if must_start else pattern
        match = re.search(regex, text)
        if not match:
            continue
        year = match.group(1)
        month = match.group(2).zfill(2)
        day = match.group(3).zfill(2) if has_day else None
        return {"year": year, "month": month, "day": day, "raw": match.group(0)}
    return None


def standardized_markdown_stem(source_path: Path, now: datetime | None = None) -> str:
    """Return JS-compatible standardized output stem for local document imports."""
    old_name = source_path.stem
    prefix_date = extract_filename_date(old_name, must_start=True, now=now)
    used_date = prefix_date or extract_filename_date(old_name, must_start=False, now=now)

    if used_date:
        year = str(used_date["year"])
        month = str(used_date["month"])
        day = used_date.get("day")
    else:
        mtime = datetime.fromtimestamp(source_path.stat().st_mtime)
        year = str(mtime.year)
        month = f"{mtime.month:02d}"
        day = f"{mtime.day:02d}"

    prefix = f"{year}{month}{day}" if day is not None else f"{year}{month}"
    if old_name.startswith(prefix + "_"):
        return safe_filename(old_name)

    rest_name = old_name
    if used_date:
        raw = str(used_date["raw"] or "")
        if prefix_date:
            rest_name = old_name[len(raw) :]
            rest_name = re.sub(r"^[-_ .]+", "", rest_name)
        else:
            rest_name = old_name.replace(raw, "", 1)
            rest_name = re.sub(r"[-_ .]+$", "", rest_name)
            rest_name = re.sub(r"\s{2,}", " ", rest_name).strip()
    rest_name = safe_filename(rest_name, default=source_path.stem)
    return safe_filename(f"{prefix}_{rest_name}", default=prefix)


def unique_path(directory: Path, filename: str) -> Path:
    candidate = directory / filename
    if not candidate.exists():
        return candidate
    stem = candidate.stem
    suffix = candidate.suffix
    for index in range(1, 1000):
        numbered = directory / f"{stem} ({index}){suffix}"
        if not numbered.exists():
            return numbered
    raise WordConvertError(f"无法生成不重复文件名：{candidate}")


def normalize_markdown(markdown: str) -> str:
    lines = [line.rstrip() for line in markdown.replace("\r\n", "\n").split("\n")]
    compacted: list[str] = []
    blank_count = 0
    for line in lines:
        if line.strip():
            compacted.append(line)
            blank_count = 0
            continue
        blank_count += 1
        if blank_count <= 2:
            compacted.append("")
    return "\n".join(compacted).strip()


def parse_document_path_inputs(values: Iterable[str]) -> list[Path]:
    """Extract unique supported document paths from CLI/menu inputs."""
    paths: list[Path] = []
    seen: set[str] = set()
    for value in values:
        candidates: list[str] = []
        text = str(value or "").strip()
        if not text:
            continue
        quoted = re.findall(rf'"([^"]+\.{SUPPORTED_SUFFIX_PATTERN})"|\'([^\']+\.{SUPPORTED_SUFFIX_PATTERN})\'', text, flags=re.I)
        candidates.extend(first or second for first, second in quoted)
        if re.match(r"^[A-Za-z]:\\", text) and Path(text.strip("\"'")).suffix.lower() in SUPPORTED_SUFFIXES:
            candidates.append(text.strip("\"'"))
        candidates.extend(re.findall(rf"[A-Za-z]:\\[^\r\n]+?\.{SUPPORTED_SUFFIX_PATTERN}", text, flags=re.I))

        for candidate in candidates:
            path = Path(candidate.strip().strip("\"'"))
            key = str(path).casefold()
            if path.suffix.lower() in SUPPORTED_SUFFIXES and key not in seen:
                seen.add(key)
                paths.append(path)
    return paths


def parse_docx_path_inputs(values: Iterable[str]) -> list[Path]:
    """Backward-compatible alias for the old Word-only parser name."""
    return parse_document_path_inputs(values)


def _qn(tag: str) -> str:
    from docx.oxml.ns import qn

    return qn(tag)


def _child_value(element, tag: str) -> str | None:
    if element is None:
        return None
    child = element.find(_qn(tag))
    if child is None:
        return None
    return child.get(_qn("w:val"))


def paragraph_num_info(paragraph) -> tuple[str | None, int]:
    p_pr = paragraph._p.pPr
    num_pr = p_pr.numPr if p_pr is not None and p_pr.numPr is not None else None
    if num_pr is None:
        return None, 0
    num_id = _child_value(num_pr, "w:numId")
    ilvl = _child_value(num_pr, "w:ilvl")
    try:
        level = int(ilvl or 0)
    except ValueError:
        level = 0
    return num_id, level


def numbering_format_map(document) -> dict[tuple[str, int], str]:
    """Return mapping from (numId, level) to Word numbering format."""
    mapping: dict[tuple[str, int], str] = {}
    try:
        numbering = document.part.numbering_part.element
    except Exception:
        return mapping

    abstract_formats: dict[tuple[str, int], str] = {}
    for abstract in numbering.findall(_qn("w:abstractNum")):
        abstract_id = abstract.get(_qn("w:abstractNumId"))
        for lvl in abstract.findall(_qn("w:lvl")):
            ilvl = lvl.get(_qn("w:ilvl")) or "0"
            fmt = _child_value(lvl, "w:numFmt") or "bullet"
            try:
                abstract_formats[(abstract_id, int(ilvl))] = fmt
            except ValueError:
                abstract_formats[(abstract_id, 0)] = fmt

    for num in numbering.findall(_qn("w:num")):
        num_id = num.get(_qn("w:numId"))
        abstract_id = _child_value(num, "w:abstractNumId")
        if num_id is None or abstract_id is None:
            continue
        for (candidate_id, level), fmt in abstract_formats.items():
            if candidate_id == abstract_id:
                mapping[(num_id, level)] = fmt
    return mapping


def inherited_run_bool(run, attr: str) -> bool:
    value = getattr(run.bold if attr == "bold" else run.italic if attr == "italic" else run.font.strike, "real", None)
    if value is not None:
        return bool(value)
    direct = run.bold if attr == "bold" else run.italic if attr == "italic" else run.font.strike
    if direct is not None:
        return bool(direct)
    return False


def run_color_hex(run) -> str | None:
    try:
        color = run.font.color
        if color is not None and color.rgb is not None:
            return str(color.rgb)
    except Exception:
        return None
    return None


def escape_markdown_text(text: str) -> str:
    return text.replace("\n", "  \n")


def render_run(run) -> str:
    text = escape_markdown_text(run.text)
    if not text:
        return ""
    rendered = text
    if inherited_run_bool(run, "bold"):
        rendered = f"**{rendered}**"
    if inherited_run_bool(run, "italic"):
        rendered = f"*{rendered}*"
    if inherited_run_bool(run, "strike"):
        rendered = f"~~{rendered}~~"
    color = run_color_hex(run)
    if color:
        rendered = f'<span style="color:#{color}">{escape(rendered, quote=False)}</span>'
    return rendered


def paragraph_text_markdown(paragraph) -> str:
    parts = [render_run(run) for run in paragraph.runs]
    text = "".join(parts).strip()
    if text:
        return text
    return paragraph.text.strip()


def paragraph_heading_level(paragraph) -> int | None:
    style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
    match = re.search(r"heading\s+([1-6])|标题\s*([1-6])", style_name, flags=re.I)
    if not match:
        return None
    return int(match.group(1) or match.group(2))


def paragraph_list_level_from_style(paragraph) -> int:
    style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
    match = re.search(r"(?:list|列表).*(\d+)$", style_name, flags=re.I)
    if match:
        return max(int(match.group(1)) - 1, 0)
    return 0


def paragraph_looks_bullet(paragraph) -> bool:
    style_name = ((paragraph.style.name if paragraph.style is not None else "") or "").lower()
    return "bullet" in style_name or "项目符号" in style_name or "列表" in style_name


def paragraph_looks_ordered(paragraph) -> bool:
    style_name = ((paragraph.style.name if paragraph.style is not None else "") or "").lower()
    return "number" in style_name or "编号" in style_name


def render_table(table) -> list[str]:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [normalize_markdown(cell.text).replace("\n", "<br>") for cell in row.cells]
        rows.append(cells)
    if not rows:
        return []
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    header = padded[0]
    lines = [
        "| " + " | ".join(cell or " " for cell in header) + " |",
        "| " + " | ".join("---" for _ in range(width)) + " |",
    ]
    for row in padded[1:]:
        lines.append("| " + " | ".join(cell or " " for cell in row) + " |")
    return lines + [""]


def iter_block_items(document):
    from docx.document import Document as _Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent = document.element.body if isinstance(document, _Document) else document._element
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def docx_to_markdown(path: Path) -> WordDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise WordConvertError("缺少或损坏依赖 python-docx/lxml，请先运行：pip install --upgrade python-docx lxml") from exc

    if not path.exists():
        raise WordConvertError(f"文件不存在：{path}")
    if path.suffix.lower() != ".docx":
        raise WordConvertError(f"暂只支持 .docx 文件：{path}")

    document = Document(str(path))
    title = (document.core_properties.title or "").strip() or path.stem
    num_formats = numbering_format_map(document)
    list_counters: dict[tuple[str, int], int] = defaultdict(int)

    lines: list[str] = []
    for block in iter_block_items(document):
        if block.__class__.__name__ == "Table":
            lines.extend(render_table(block))
            continue

        paragraph = block
        text = paragraph_text_markdown(paragraph)
        if not text:
            if lines and lines[-1] != "":
                lines.append("")
            continue

        heading_level = paragraph_heading_level(paragraph)
        if heading_level:
            lines.extend([f"{'#' * heading_level} {text}", ""])
            continue

        num_id, level = paragraph_num_info(paragraph)
        fmt = num_formats.get((num_id, level)) if num_id is not None else None
        style_level = paragraph_list_level_from_style(paragraph)
        if num_id is not None or paragraph_looks_bullet(paragraph) or paragraph_looks_ordered(paragraph):
            level = level if num_id is not None else style_level
            indent = "  " * max(level, 0)
            if fmt and fmt != "bullet" or (fmt is None and paragraph_looks_ordered(paragraph)):
                key = (num_id or "style-numbered", level)
                list_counters[key] += 1
                marker = f"{list_counters[key]}. "
            else:
                marker = "- "
            lines.append(f"{indent}{marker}{text}")
            continue

        lines.extend([text, ""])

    markdown = normalize_markdown("\n".join(lines))
    if not markdown:
        raise WordConvertError(f"没有从 Word 文档中解析出正文：{path}")
    return WordDocument(title=title, markdown=markdown, source_path=path)


def find_office_converter() -> str | None:
    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\Program Files\OpenOffice 4\program\soffice.exe",
        r"C:\Program Files (x86)\OpenOffice 4\program\soffice.exe",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


def convert_doc_to_docx(path: Path, temp_dir: Path) -> Path:
    converter = find_office_converter()
    converted = temp_dir / f"{path.stem}.docx"
    if converter:
        cmd = [
            converter,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(temp_dir),
            str(path),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
        if result.returncode == 0 and converted.exists():
            return converted
        detail = (result.stderr or result.stdout or "").strip()
        raise WordConvertError(f".doc 转 .docx 失败：{detail or '未知错误'}")

    if sys.platform == "win32":
        try:
            import win32com.client  # type: ignore[import-not-found]
        except ImportError:
            pass
        else:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(str(path.resolve()))
                doc.SaveAs2(str(converted.resolve()), FileFormat=16)
                doc.Close(False)
            finally:
                word.Quit()
            if converted.exists():
                return converted

    raise WordConvertError(
        "暂无法转换 .doc 文件：需要安装 LibreOffice/OpenOffice，或安装 Microsoft Word 与 pywin32。"
    )


def doc_to_markdown(path: Path) -> WordDocument:
    if not path.exists():
        raise WordConvertError(f"文件不存在：{path}")
    with tempfile.TemporaryDirectory(prefix="doc_to_md_") as temp:
        converted = convert_doc_to_docx(path, Path(temp))
        document = docx_to_markdown(converted)
    document.title = path.stem
    document.source_path = path
    return document


def pdf_to_markdown(path: Path) -> WordDocument:
    if not path.exists():
        raise WordConvertError(f"文件不存在：{path}")
    try:
        import pdfplumber
    except ImportError as exc:
        raise WordConvertError("缺少依赖 pdfplumber，请先运行：pip install pdfplumber") from exc

    pages: list[str] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            title = ((pdf.metadata or {}).get("Title") or "").strip() or path.stem
            for index, page in enumerate(pdf.pages, 1):
                text = page.extract_text(layout=True) or page.extract_text() or ""
                text = normalize_markdown(text)
                if text:
                    pages.extend([f"<!-- PDF 第 {index} 页 -->", "", text, ""])
    except Exception as exc:
        raise WordConvertError(f"PDF 解析失败：{exc}") from exc

    markdown = normalize_markdown("\n".join(pages))
    if not markdown:
        raise WordConvertError(f"没有从 PDF 中解析出文字：{path}。如果是扫描版 PDF，需要先 OCR。")
    return WordDocument(title=title, markdown=markdown, source_path=path)


def read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    encodings = ("utf-8-sig", "utf-8", "gb18030", "cp936", "big5")
    for encoding in encodings:
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def txt_to_markdown(path: Path) -> WordDocument:
    if not path.exists():
        raise WordConvertError(f"文件不存在：{path}")
    text = normalize_markdown(read_text_file(path))
    if not text:
        raise WordConvertError(f"TXT 文件为空或没有可解析文字：{path}")
    return WordDocument(title=path.stem, markdown=text, source_path=path)


def document_to_markdown(path: Path) -> WordDocument:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return docx_to_markdown(path)
    if suffix == ".doc":
        return doc_to_markdown(path)
    if suffix == ".pdf":
        return pdf_to_markdown(path)
    if suffix == ".txt":
        return txt_to_markdown(path)
    raise WordConvertError(f"暂不支持的文件类型：{path.suffix or '(无扩展名)'}")


def source_uri(path: Path) -> str:
    try:
        return path.resolve().as_uri()
    except Exception:
        return str(path)


def build_word_markdown(document: WordDocument, body_markdown: str | None = None) -> str:
    body = body_markdown if body_markdown is not None else document.markdown
    body = bump_headings_when_h1_exists(normalize_markdown(body))
    lines = [
        "# 基本信息",
        "",
        f"- **日期**: {datetime.now().strftime('%Y-%m-%d')}",
        f"- **原文链接**: [{document.title}]({source_uri(document.source_path)})",
        "- **行业**: ",
        "- **公司**: ",
        "- **人工标签**: ",
        "- [ ] **是否已读**",
        "- **我的评价**: ",
        "",
        "---",
        "",
        "# 正文",
        "",
        body,
        "",
    ]
    return "\n".join(lines)


def convert_word_doc(
    path: str | Path,
    out_dir: str | Path = DEFAULT_OUT_DIR,
    *,
    overwrite: bool = False,
) -> Path:
    document = document_to_markdown(Path(path))
    output_dir = Path(out_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = standardized_markdown_stem(document.source_path) + ".md"
    output_path = output_dir / filename if overwrite else unique_path(output_dir, filename)
    output_path.write_text(normalize_markdown_output(build_word_markdown(document)), encoding="utf-8")
    return output_path


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="把本地 docx/doc/pdf/txt 文档转换为会议纪要 Inbox Markdown。")
    parser.add_argument("paths", nargs="+", help="一个或多个本地 docx/doc/pdf/txt 文件路径")
    parser.add_argument(
        "--out-dir",
        default=str(DEFAULT_OUT_DIR),
        help=f"Markdown 保存目录，默认：{DEFAULT_OUT_DIR}",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="如果同名 Markdown 已存在，直接覆盖而不是生成 (1)、(2) 文件。",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    paths = parse_document_path_inputs(args.paths)
    if not paths:
        print("没有解析到有效的 docx/doc/pdf/txt 文件路径。", file=sys.stderr)
        return 2

    saved_paths: list[Path] = []
    failed: list[tuple[Path, str]] = []
    for index, path in enumerate(paths, 1):
        print(f"[{index}/{len(paths)}] 转换: {path}")
        try:
            output_path = convert_word_doc(path, out_dir=Path(args.out_dir), overwrite=args.overwrite)
            saved_paths.append(output_path)
            print(f"  已保存：{output_path}")
        except WordConvertError as exc:
            failed.append((path, str(exc)))
            print(f"  失败：{exc}", file=sys.stderr)
        except Exception as exc:
            failed.append((path, str(exc)))
            print(f"  转换异常：{exc}", file=sys.stderr)

    print()
    print(f"完成：成功 {len(saved_paths)} 篇，失败 {len(failed)} 篇。")
    if failed:
        print("失败文件：", file=sys.stderr)
        for path, reason in failed:
            print(f"- {path} -> {reason}", file=sys.stderr)
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
